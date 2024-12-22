from docx import Document
from docx.shared import Inches
#from fpdf import FPDF
#import os
import streamlit as st
from lxml import etree
from datetime import datetime

def format_number(value):
    """Format numbers as 'xxx,xxx.xx'."""
    return f"{value:,.2f}"

# Function to replace placeholders in text boxes
def replace_text_in_text_boxes(doc, placeholders):
    xml_tree = etree.fromstring(doc._element.xml)  # Parse the XML
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    for placeholder, value in placeholders.items():
        for elem in xml_tree.xpath(f".//w:txbxContent//w:t[text()[contains(.,'{placeholder}')]]", namespaces=namespaces):
            elem.text = elem.text.replace(placeholder, value)

    # Update the document's XML
    doc._element = etree.ElementTree(xml_tree).getroot()

# Function to replace an image placeholder
def replace_image(doc, image_path, placeholder="image_placeholder"):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.clear()  # Clear the text placeholder
                # Insert the image in place of the placeholder
                run.add_picture(image_path, width=Inches(2))  # Adjust size as needed

# Function to generate DOCX with placeholder replacements
def generate_docx(image_file=None):
    template_path = "Intial.docx"
    output_path = "output.docx"
    doc = Document(template_path)

    # Replace placeholders in regular paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in placeholders.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in placeholders.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)

    # Replace placeholders in text boxes
    replace_text_in_text_boxes(doc, placeholders)

    # Replace image if uploaded
    if image_file is not None:
        replace_image(doc, image_file)

    # Save the modified DOCX file
    doc.save(output_path)
    return output_path

# Streamlit code for input and calculation logic
st.title('Initial Proposal Creator')

Quote = st.text_input("Quote reference:", value=None)
Site_Name = st.text_input("Site Name:", value=None)
Address = st.text_input("Site Address:", value=None)
City = st.text_input("Site City:", value=None)
Post_code = st.text_input("Site PostCode:", value=None)
Pdate = st.date_input("Proposal Date", value=datetime.now())
annual_consumption = st.number_input("Annual Consumption in kWh:", format="%.2f", value=0.0)
estimated_percentage = st.number_input("Estimated Saving Percentage:", format="%.2f")
tarif = st.number_input("Tariff £/kWh:", value=0.0, format="%.2f")
vo_size = st.selectbox("VO size", ('63', '100', '125', '160', '200', '250', '315', '400', '500', '630', '800', '1000', '1250', '1600', '2250'))

capital_cost_predefined = {
    '63': "5645",
    '100': "6253",
    '125': "9383",
    '160': "11004",
    '200': "12468",
    '250': "15023",
    '315': "17174",
    '400': "19177",
    '500': "24241",
    '630': "30004",
    '800': "37706",
    '1000': "47166",
    '1250': "55440",
    '1600': "60085",
    '2250': "74545"
}

capital_cost_default = capital_cost_predefined[vo_size]

toggle = st.checkbox("Override capital cost")
if toggle:
    capital_cost = st.text_input(f"Enter value for {vo_size}a:", capital_cost_default)
else:
    capital_cost = st.text_input(f"Total Cost + Installation for {vo_size}a:", capital_cost_default, disabled=True)

# Image upload for replacement
uploaded_image = st.file_uploader("Upload an image for the proposal", type=["jpg", "png", "jpeg"])

def calculate():
    annual_spend = annual_consumption * tarif
    kWh_saving = round((estimated_percentage / 100) * annual_consumption)
    cost_saving = round(annual_spend * (estimated_percentage / 100))
    co2_savings = round((kWh_saving * 0.207074)/1000)
    ROI_in_years = round(float(capital_cost) / cost_saving)
    ROI_in_months = round(ROI_in_years * 12)
    kWh_saving_5yr = round(kWh_saving * 5)
    kWh_saving_10yr = round(kWh_saving * 10)
    cost_saving_5yr = round(cost_saving * 5)
    cost_saving_10yr = round(cost_saving * 10)
    co2_savings_5yr = round(co2_savings * 5)
    co2_savings_10yr = round(co2_savings * 10)
    return annual_spend, kWh_saving, cost_saving, co2_savings, ROI_in_years, ROI_in_months, kWh_saving_5yr, kWh_saving_10yr, cost_saving_5yr, cost_saving_10yr, co2_savings_5yr, co2_savings_10yr

if st.button("Generate and Download DOCX"):
    try:
        annual_spend, kWh_saving, cost_saving, co2_savings, ROI_in_years, ROI_in_months, kWh_saving_5yr, kWh_saving_10yr, cost_saving_5yr, cost_saving_10yr, co2_savings_5yr, co2_savings_10yr = calculate()
        placeholders = {
            "{{Quote}}": Quote,
            "{{Sitename}}" : Site_Name,
            "{{Address}}" : Address,
            "{{City}}" : City,
            "{{Postcode}}" : Post_code,
            "{{date}}" : Pdate.strftime("%d/%m/%Y"),  # Format date
            "{{Annual spend}}" : f"{format_number(annual_spend)}",
            "{{Annual Consumption kWh}}": f"{format_number(annual_consumption)}",
            "{{tariff}}": f"{tarif}",
            "{{VO Size}}": vo_size,
            "{{cost1}}": f"{format_number(cost_saving)}",
            "{{Capital Cost}}": f"{format_number(float(capital_cost))}",
            "{{ROI Months}}": f"{ROI_in_months}",
            "{{kWh}}": f"{format_number(kWh_saving)}",
            "{{per}}": f"{estimated_percentage}",
            "{{co21}}": f"{format_number(co2_savings)}",
            "{{kWh5}}": f"{format_number(kWh_saving_5yr)}",
            "{{kWh10}}": f"{format_number(kWh_saving_10yr)}",
            "{{cost5}}": f"{format_number(cost_saving_5yr)}",
            "{{cost10}}": f"{format_number(cost_saving_10yr)}",
            "{{co25}}": f"{format_number(co2_savings_5yr)}",
            "{{co210}}": f"{format_number(co2_savings_10yr)}"
        }

        docx_path = generate_docx(uploaded_image)

        with open(docx_path, "rb") as f:
            st.download_button(label="Download DOCX", data=f, file_name="Proposal.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.error(f"Error generating DOCX: {e}")