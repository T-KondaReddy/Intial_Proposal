import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import os
from lxml import etree
from docx.shared import Inches
from docx import Document
from datetime import datetime

# Set the page configuration
st.set_page_config(layout="wide")


# Function to replace placeholders in text boxes
def replace_text_in_text_boxes(doc, placeholders):
    xml_tree = etree.fromstring(doc._element.xml)  # Parse the XML
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    for placeholder, value in placeholders.items():
        for elem in xml_tree.xpath(f".//w:txbxContent//w:t[text()[contains(.,'{placeholder}')]]", namespaces=namespaces):
            elem.text = elem.text.replace(placeholder, value)

    # Update the document's XML
    doc._element = etree.ElementTree(xml_tree).getroot()

# Function to replace an image placeholder
def replace_image(doc, image_path, placeholder, inches, hinches):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.clear()  # Clear the text placeholder
                # Insert the image in place of the placeholder
                run.add_picture(image_path, width=Inches(inches), height= Inches(hinches))  # Adjust size as needed

# Function to generate DOCX with placeholder replacements
def generate_docx(image_file=None, creator=None):
    if creator:
        template_path = "./Final_Proposal/CVO_Final.docx"
    else:
        template_path = "./Final_Proposal/final.docx"
    output_path = "./Final_Proposal/final_output.docx"
    doc = Document(template_path)

    # Replace placeholders in regular paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in placeholders.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in placeholders.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)

    # Replace placeholders in text boxes
    replace_text_in_text_boxes(doc, placeholders)

    # Replace image if uploaded
    if image_file is not None:
        placeholder = "image_placeholder"
        replace_image(doc, image_file,placeholder, 3, 2)
    
    # Replace image if uploaded
    if image_file is not None:
        placeholder1 = "image_placeholder1"
        replace_image(doc, image_file,placeholder1, 0.5,0.5)

    graph1 = None
    graph1="./graphs/yearly_savings_stylized.jpg"
    placeholder_graph1 = "graph1"

    if graph1 is not None:
        replace_image(doc, graph1,placeholder_graph1, 4,2.5)

    graph2 = None
    graph2="./graphs/cumulative_vs_capital_stylized.jpg"
    placeholder_graph2 = "graph2"

    if graph2 is not None:
        replace_image(doc, graph2,placeholder_graph2, 4,2.5)

    # Save the modified DOCX file
    doc.save(output_path)
    return output_path


# Page title
st.title("VMS Saving")

# Voltage to savings percentage mapping
savings_map = {
    225: [6, 0, 3, 0, 3, 9, 1, 2, 0, 0],
    230: [11, 1, 6, 0, 4, 13, 1, 3, 0, 0],
    235: [18, 1, 11, 1, 4, 16, 2, 4, 1, 0],
    240: [20, 1, 12, 1, 6, 16, 2, 5, 1, 0],
    245: [20, 1, 14, 2, 7, 16, 3, 5, 2, 0],
    250: [22, 1, 16, 3, 8, 18, 3, 6, 3, 0],
    255: [22, 1, 16, 3, 8, 18, 3, 7, 3, 0],
}

# Create the data for the table
data = {
    "Load Profile": [
        "Lighting with Magnetic Ballast (Old Style Lighting with T10/T8 Tube etc)",
        "Lighting with Electronic Control (T5/LED/Electronic Ballast/High Frequency)",
        "Refrigeration/Chillers/Freezers/Air Conditioning (Without Inverter)",
        "Refrigeration/Chillers/Freezers/Air Conditioning (With Inverter)",
        "Heating Ventilation and Air Conditioning (HVAC)",
        "Heater/Kitchen Equipment/Resistive Loads (Without Thermocouple)",
        "Computer/Electronic/General IT Equipment",
        "Motors/Machinery/Pumps etc",
        "Inverter Drives",
        "Other Non Saving Load",
    ],
    "Load %": [0] * 10,  # Initially all zeros
    "Saving %": [0] * 10,  # Initially all zeros
    "Total Saving": [0.0] * 10,  # Initially all zeros
    "Identifier": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10],
}

df = pd.DataFrame(data)

# Value section
Value_col1, Value_col2, Value_col3 = st.columns([1, 1, 1])
with Value_col1:
    psr_type = st.selectbox("Select PSR Type:", ["Type S", "Hybrid Type H", "Compact Type C"], key="psr_type")
with Value_col2:
    # Input for average voltage
    average_voltage = st.number_input("Enter Average Voltage:", min_value=225, max_value=255, step=5, value=240)
with Value_col3:
    tariff = st.number_input("Tariff (£/kWh):", value=0.0, format="%.2f")

# Update the Savings % based on the average voltage
if average_voltage in savings_map:
    df["Saving %"] = savings_map[average_voltage]

# Calculate Total Saving based on PSR Type
def calculate_total_saving(load_percent, saving_percent, psr_type):
    # Base calculation: Load % * Saving % 
    base_saving = load_percent * saving_percent / 100
    
    # Adjust based on PSR type
    if psr_type == "Type S":
        return base_saving
    elif psr_type == "Hybrid Type H":
        return base_saving * 0.9
    elif psr_type == "Compact Type C":
        return base_saving * 0.8
    return base_saving  # Default to no adjustment

col2, col3, col4, col5 = st.columns([5, 1, 1, 1])

# Lists to store updated values
load_profile_values = []
load_percentages = []
saving_percentages = []
total_savings = []

with col2:
    st.markdown("**Load Profile**")
    for i, item in enumerate(df["Load Profile"]):
        st.text_input("No Label", value=item, disabled=True, key=f"profile_{i}", label_visibility="collapsed")
        load_profile_values.append(item)  # Store Load Profile values

with col3:
    st.markdown("**Load %**")
    for i, value in enumerate(df["Load %"]):
        updated_value = st.number_input(
            "No Label", value=value, key=f"load_{i}", label_visibility="collapsed"
        )
        df.at[i, "Load %"] = updated_value  # Update dataframe
        load_percentages.append(updated_value)  # Store Load % values

with col4:
    st.markdown("**Saving %**")
    for i, value in enumerate(df["Saving %"]):
        st.number_input(
            "No Label", value=value, key=f"saving_{i}", disabled=True, label_visibility="collapsed"
        )
        saving_percentages.append(value)  # Store Saving % values

with col5:
    st.markdown("**Total Saving**")
    for i in range(len(df)):
        # Calculate Total Saving based on PSR Type
        total_saving = float(calculate_total_saving(df.at[i, "Load %"], df.at[i, "Saving %"], psr_type))
        df.at[i, "Total Saving"] = round(total_saving, 2)  # Update dataframe
        st.number_input(
            "No Label",
            value=df.at[i, "Total Saving"],
            key=f"total_saving_{i}",
            disabled=True,
            label_visibility="collapsed",
        )
        total_savings.append(df.at[i, "Total Saving"])  # Store Total Saving values

# Now you have all the values stored in variables:
# - `load_profile_values`: List of Load Profile values
# - `load_percentages`: List of updated Load % values
# - `saving_percentages`: List of Saving % values
# - `total_savings`: List of calculated Total Saving values


# Total Saving % calculation (sum of Total Savings as percentage)
total_saving_sum = df["Total Saving"].sum()

# Add a toggle button to override the Total Saving %
override_total_saving = st.checkbox("Override Total Saving %")


# New section for energy data
energy_col1, energy_col2, energy_col3 = st.columns([1, 1, 1])
with energy_col1:
    if override_total_saving:
        total_saving_override = st.number_input("Override Total Saving %:", value=total_saving_sum, format="%.2f")
    else:
        total_saving_override = st.number_input("Override Total Saving %:", value=total_saving_sum, format="%.2f", disabled=True)
with energy_col2:
    total_kWh_per_annum = st.number_input("Total kWh Per Annum", value=0.00, format="%.2f")
with energy_col3:
    co2_conversion = st.number_input("T CO2 - kWh Conversion (T/kWh)", value=0.000233, format="%.6f", disabled=True)

energy_col4, energy_col5, energy_col6 = st.columns([1, 1, 1])
with energy_col4:
    ccl_charge = st.number_input("£ / CCL", value=0.0078, format="%.4f", disabled=True)
with energy_col5:
    ccl_percent = st.number_input("CCL % Charge", value=100, disabled=True)
with energy_col6:
    energy_price_increase = st.number_input("Annual Energy Price Increase (%)", value=8.00 , format="%.2f", disabled=True)

# Calculate new values based on the formulas
kWh_saving = total_kWh_per_annum * (total_saving_override / 100)
co2_saving = kWh_saving * co2_conversion
financial_kWh_savings = kWh_saving * tariff
financial_kWh_ccl = kWh_saving * (ccl_percent / 100) * ccl_charge

# Display the results in a non-editable text field
result_col1, result_col2, result_col3, result_col4 = st.columns([1, 1, 1, 1])

with result_col1:
    st.text_input("kWh Saving (kWh)", value=f"{kWh_saving:.4f}", disabled=True)
with result_col2:
    st.text_input("CO2 Saving (Tons)", value=f"{co2_saving:.2f}", disabled=True)
with result_col3:
    st.text_input("Financial kWh Savings (£)", value=f"£{financial_kWh_savings:.2f}", disabled=True)
with result_col4:
    st.text_input("Financial kWh CCL (£)", value=f"£{financial_kWh_ccl:.2f}", disabled=True)

total_annual_saving = financial_kWh_savings + financial_kWh_ccl

total_col1, total_col2 = st.columns([1, 1])
with total_col1:
    st.text_input("Total Annual Saving(£)", value=f"£{total_annual_saving:.2f}", disabled=True)

        
    
# VO Size and costs
vo_size = st.selectbox("VO size", ('63', '100', '125', '160', '200', '250', '315', '400', '500', '630', '800', '1000', '1250', '1600', '258/170', '380/250', '523/350', '653/430', '805/530', '1064/700', '1519/1000', '2157/1400', '3191/2100', '4862/3200'))

# New cost data
capital_cost_data = {
    '63': (2409, 1630, 600),
    '100': (2774, 1630, 600),
    '125': (4652, 1630, 600),
    '160': (5025, 2630, 600),
    '200': (5903, 2630, 600),
    '250': (7436, 2630, 600),
    '315': (7929, 3959, 600),
    '400': (9131, 3959, 600),
    '500': (12169, 3959, 600),
    '630': (15627, 3959, 600),
    '800': (19481, 5238, 600),
    '1000': (25157, 5238, 600),
    '1250': (30121, 5238, 600),
    '1600': (32908, 5238, 600),
    '258/170': (14095, 3458, 650),
    '380/250': (16169, 3458, 650),
    '523/350': (20207, 3956, 650),
    '653/430': (23466, 3956, 650),
    '805/530': (26862, 4405, 650),
    '1064/700': (32652, 5235, 650),
    '1519/1000': (48312, 5235, 650),
    '2157/1400': (55297, 6315, 650),
    '3191/2100': (82207, 6315, 650),
    '4862/3200': (78255, 7395, 650),
}

# Get the cost for selected VO size
total_cost, installation_cost,meter_cost = capital_cost_data.get(vo_size, (0, 0, 0))
Non_Smartmeter = st.checkbox("Non Smart Meter (NM)")
if Non_Smartmeter:
    total_cost = total_cost-meter_cost
# Display Total cost and Installation cost
col1, col2 = st.columns([2, 2])
with col1:
    st.number_input("Total Cost (£)", value=total_cost, disabled=True)
with col2:
    st.number_input("Installation Cost (£)", value=installation_cost, disabled=True)

final_col1,final_col2 = st.columns([1,1])

with final_col1:
# Margin percentage input
    margin_percent = st.number_input("Enter Margin Percentage (%)", min_value=0.0, max_value=100.0, value=40.0)

with final_col2:
# Miscellaneous cost input
    misc_cost = st.number_input("Additional Miscellaneous Cost (£)", value=0.0, format="%.2f")
cost_col1,cost_col2 = st.columns([1,1])
with cost_col1:
# Checkbox to add £1000 to capital cost
    add_1000_checkbox = st.checkbox("Add lift & shift")

# Calculate the capital cost after margin and add-ons
capital_cost_with_margin = total_cost /((100 - margin_percent) / 100) + installation_cost + misc_cost
if add_1000_checkbox:
    capital_cost_with_margin += 1000

# Display the capital cost field and handle override
capital_cost_final = int(capital_cost_with_margin)

override_capital_cost = st.checkbox("Override Capital Cost (£)", value=False)

if override_capital_cost:
    capital_cost_final = st.number_input("Override Capital Cost (£)", value=capital_cost_with_margin, format="%.2f")
else:
    st.text_input("Capital Cost (£)", value=f"£{capital_cost_final:.2f}", disabled=True)

# Calculate yearly savings with 6% increase and cumulative savings
years = 10
yearly_savings = []
cumulative_savings = []

current_saving = total_annual_saving
cumulative_total = 0

for year in range(1, years + 1):
    yearly_savings.append(round(current_saving, 2))
    cumulative_total += current_saving
    cumulative_savings.append(round(cumulative_total, 2))
    current_saving = current_saving * 1.08  # Increase by 6%

# Create a DataFrame for the table
savings_table = pd.DataFrame({
    "Year": list(range(1, years + 1)),
    "Yearly Saving (£)": yearly_savings,
    "Cumulative Saving (£)": cumulative_savings
})

# Display the table at the end of the Streamlit app
#st.markdown("### 10-Year Savings Projection")
#st.table(savings_table)

# Input fields for 5th and 10th Year Cumulative Saving
cumulative_5th_year = cumulative_savings[4]  # 5th year (index 4)
cumulative_10th_year = cumulative_savings[9]  # 10th year (index 9)

col1, col2 = st.columns([1, 1])
with col1:
    st.text_input("Cumulative Saving (5th Year) (£)", value=f"£{cumulative_5th_year:.2f}", disabled=True)
with col2:
    st.text_input("Cumulative Saving (10th Year) (£)", value=f"£{cumulative_10th_year:.2f}", disabled=True)

# Function to format numbers as "xxx.xxx,xx"
def format_number(value):
    return f"{value:,.2f}"

# Set folder path to save graphs
output_folder = "graphs"
os.makedirs(output_folder, exist_ok=True)

# Columns for side-by-side graphs
graph_col1, graph_col2 = st.columns(2)

# Bar Chart for Yearly Savings
with graph_col1:
    fig_bar, ax_bar = plt.subplots(figsize=(4, 3))  # Adjust size as required

    # Create bar chart
    bars = ax_bar.bar(range(1, years + 1), yearly_savings, color="#0070C0", edgecolor="black", alpha=0.8)

    # Add £ values on top of each bar
    for bar in bars:
        height = bar.get_height()
        ax_bar.text(bar.get_x() + bar.get_width() / 2, height, f"£{int(round(height))}", ha="center", va="bottom", fontsize=5)

    # Formatting
    ax_bar.set_title("VMS Financial Savings Per Annum", fontsize=10, fontweight="bold", color="black")
    ax_bar.set_xlabel("Year", fontsize=8, color="black")
    ax_bar.set_ylabel("Savings (£)", fontsize=8, color="black")
    ax_bar.set_xticks(range(1, years + 1))
    ax_bar.set_xticklabels(range(1, years + 1), fontsize=7, rotation=0)
    ax_bar.tick_params(axis="y", labelsize=7)
    #ax_bar.grid(axis="y", linestyle="--", linewidth=0.5, alpha=0.7)  # Add gridlines
    plt.tight_layout()

    # Save bar chart as JPG
    bar_chart_path = os.path.join(output_folder, "yearly_savings_stylized.jpg")
    plt.savefig(bar_chart_path, format="jpg", dpi=300)
    st.pyplot(fig_bar)

# Line Graph for Cumulative Savings vs Capital Cost
with graph_col2:
    fig_line, ax_line = plt.subplots(figsize=(4, 3))  # Adjust size as required

    # Plot cumulative savings in green solid line
    ax_line.plot(range(1, years + 1), cumulative_savings, label="Cumulative Savings (£)", color="green", linestyle="-", marker="o", linewidth=1, markersize=3, alpha=0.8)
    
    # Plot capital cost in blue solid line
    ax_line.plot(range(1, years + 1), [capital_cost_final] * years, label="Capital Cost (£)", color="#0070C0", linestyle="-", marker="*", linewidth=1, markersize=4, alpha=0.8)
    ax_line.text(10,capital_cost_final, f"£ {capital_cost_final}", ha="right", va="bottom", fontsize=5)

    # Add £ values on cumulative savings
    for i, value in enumerate(cumulative_savings):
        if i != 0:
            ax_line.text(i + 1, value, f"£{int(value)}", ha="right", va="bottom", fontsize=5)
        

    # Formatting
    ax_line.set_title("VMS Savings Return On Investment", fontsize=10, fontweight="bold", color="black")
    ax_line.set_xlabel("Year", fontsize=8, color="black")
    ax_line.set_xticks(range(1, years + 1))
    ax_line.set_xticklabels(range(1, years + 1), fontsize=7, rotation=0)
    ax_line.tick_params(axis="y", labelsize=7)
    ax_line.legend(fontsize=7, loc="upper left")
    ax_line.grid(axis="y", linestyle="dotted", linewidth=0.5, alpha=0.7)  # Add gridlines
    plt.tight_layout()

    # Save line graph as JPG
    line_chart_path = os.path.join(output_folder, "cumulative_vs_capital_stylized.jpg")
    plt.savefig(line_chart_path, format="jpg", dpi=300)
    st.pyplot(fig_line)


ROI_Years=capital_cost_final/total_annual_saving
ROI_Months = ROI_Years*12

# Confirm saved paths

roi_col1, roi_col2 = st.columns([2, 2])
with roi_col1:
    st.text_input("ROI in Years", value=f"{ROI_Years:.2f} Years", disabled=True)
with roi_col2:
    st.text_input("ROI in Months", value=f"{ROI_Months:.2f} Months", disabled=True)

# Define the layout in a 3x2 grid
col1, col2, col3 = st.columns(3)

with col1:
    Quote = st.text_input("Quote reference:", value=None)

with col2:
    Site_Name = st.text_input("Site Name:", value=None)

with col3:
    Pdate = st.date_input("Proposal Date", value=datetime.now())

about_Site = st.text_input (f"About {Site_Name} (in 50 words):")
# Image upload for replacement
uploaded_image = st.file_uploader("Upload Company logo", type=["jpg", "png", "jpeg"])

EMCs_Site = st.checkbox("Check if Paul Beck is the Proposal Creator", value=False)
if st.button("Generate and Download DOCX"):
    try:
        placeholders = {
            "{{Quote}}": Quote,
            "{{Sitename}}" : Site_Name,
            "{{about}}" : about_Site,
            "{{date}}" : Pdate.strftime("%d/%m/%Y"),  # Format date
            "{{Annual spend}}" : f"{format_number(total_annual_saving)}",
            "{{Annual Consumption kWh}}": f"{format_number(total_kWh_per_annum)}",
            "{{tariff}}": f"{tariff}",
            "{{VO Size}}": vo_size,
            "{{Capital Cost}}": f"{format_number(float(capital_cost_final))}",
            "{{ROI Months}}": f"{format_number(ROI_Months)}",
            "{{ROI Years}}" : f"{format_number(ROI_Years)}",
            "{{kWh}}": f"{format_number(kWh_saving)}",
            "{{per}}": f"{format_number(total_saving_override)}",
            "{{co21}}": f"{format_number(co2_saving)}",
            "{{cost5}}": f"{format_number(cumulative_5th_year)}",
            "{{cost10}}": f"{format_number(cumulative_10th_year)}",
            "{{£kWh}}":  f"{format_number(financial_kWh_savings)}",
            "{{CCL}}" : f"{format_number(financial_kWh_ccl)}",
            "{{sp1}}": f"{saving_percentages[0]}",
            "{{sp2}}": f"{saving_percentages[1]}",
            "{{sp3}}": f"{saving_percentages[2]}",
            "{{sp4}}": f"{saving_percentages[3]}",
            "{{sp5}}": f"{saving_percentages[4]}",
            "{{sp6}}": f"{saving_percentages[5]}",
            "{{sp7}}": f"{saving_percentages[6]}",
            "{{sp8}}": f"{saving_percentages[7]}",
            "{{sp9}}": f"{saving_percentages[8]}",
            "{{sp10}}": f"{saving_percentages[9]}",
            "{{l%1}}": f"{load_percentages[0]}",
            "{{l%2}}": f"{load_percentages[1]}",
            "{{l%3}}": f"{load_percentages[2]}",
            "{{l%4}}": f"{load_percentages[3]}",
            "{{l%5}}": f"{load_percentages[4]}",
            "{{l%6}}": f"{load_percentages[5]}",
            "{{l%7}}": f"{load_percentages[6]}",
            "{{l%8}}": f"{load_percentages[7]}",
            "{{l%9}}": f"{load_percentages[8]}",
            "{{l%10}}": f"{load_percentages[9]}",
            "{{ts1}}": f"{total_savings[0]}",
            "{{ts2}}": f"{total_savings[1]}",
            "{{ts3}}": f"{total_savings[2]}",
            "{{ts4}}": f"{total_savings[3]}",
            "{{ts5}}": f"{total_savings[4]}",
            "{{ts6}}": f"{total_savings[5]}",
            "{{ts7}}": f"{total_savings[6]}",
            "{{ts8}}": f"{total_savings[7]}",
            "{{ts9}}": f"{total_savings[8]}",
            "{{ts10}}": f"{total_savings[9]}",
            
        }
        docx_path = generate_docx(uploaded_image, EMCs_Site)
        with open(docx_path, "rb") as f:
            st.download_button(label="Download DOCX", data=f, file_name="Proposal.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.error(f"Error generating DOCX: {e}")